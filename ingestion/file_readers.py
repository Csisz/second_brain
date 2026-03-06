"""
ingestion/file_readers.py
Fájl formátum olvasók: .docx, .pdf, .xlsx, .txt, .md,
                       .py/.ps1/.sh/.bat, .csv, .xml/.json/.yaml, .html
"""
import io
from pathlib import Path


# ─── Meglévő olvasók ────────────────────────────────────────────────────────

def read_docx(path: str | Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_pdf(path: str | Path) -> str:
    import PyPDF2
    text_parts = []
    with open(str(path), "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                text_parts.append(t.strip())
    return "\n".join(text_parts)


def read_xlsx(path: str | Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) for c in row if c is not None and str(c).strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_txt(path: str | Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


# ─── Script olvasók (.py, .ps1, .sh, .bat) ──────────────────────────────────

def read_script(path: str | Path) -> str:
    content = Path(path).read_text(encoding="utf-8", errors="ignore")
    suffix = Path(path).suffix.lower()
    lang_map = {
        ".py": "Python", ".ps1": "PowerShell",
        ".sh": "Bash/Shell", ".bat": "Windows Batch", ".cmd": "Windows Batch"
    }
    lang = lang_map.get(suffix, "Script")
    return f"[{lang} script: {Path(path).name}]\n\n{content}"


# ─── CSV olvasó ──────────────────────────────────────────────────────────────

def read_csv(path: str | Path) -> str:
    import csv
    parts = []
    try:
        with open(str(path), encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)
            if not rows:
                return ""
            header = rows[0]
            parts.append("Oszlopok: " + " | ".join(header))
            for row in rows[1:501]:
                row_text = " | ".join(str(c).strip() for c in row if str(c).strip())
                if row_text:
                    parts.append(row_text)
            if len(rows) > 501:
                parts.append(f"[... összesen {len(rows)-1} sor, első 500 betöltve]")
    except Exception:
        return read_txt(path)
    return "\n".join(parts)


# ─── XML olvasó ──────────────────────────────────────────────────────────────

def read_xml(path: str | Path) -> str:
    import xml.etree.ElementTree as ET
    try:
        tree = ET.parse(str(path))
        root = tree.getroot()
        parts = [f"[XML: {Path(path).name}]"]

        def extract(element, depth=0):
            indent = "  " * depth
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            text = element.text.strip() if element.text and element.text.strip() else ""
            attribs = " ".join(f'{k}="{v}"' for k, v in element.attrib.items())
            line = f"{indent}<{tag}"
            if attribs:
                line += f" {attribs}"
            if text:
                line += f"> {text}"
            parts.append(line)
            for child in element:
                extract(child, depth + 1)

        extract(root)
        return "\n".join(parts)
    except Exception:
        return read_txt(path)


# ─── JSON olvasó ─────────────────────────────────────────────────────────────

def read_json(path: str | Path) -> str:
    import json
    try:
        content = Path(path).read_text(encoding="utf-8", errors="ignore")
        if len(content) > 50000:
            content = content[:50000] + "\n[... fájl csonkítva 50KB-nál]"
        data = json.loads(content)
        formatted = json.dumps(data, ensure_ascii=False, indent=2)
        if len(formatted) > 50000:
            formatted = formatted[:50000] + "\n[... csonkítva]"
        return f"[JSON: {Path(path).name}]\n\n{formatted}"
    except Exception:
        return read_txt(path)


# ─── YAML olvasó ─────────────────────────────────────────────────────────────

def read_yaml(path: str | Path) -> str:
    try:
        import yaml
        content = Path(path).read_text(encoding="utf-8", errors="ignore")
        data = yaml.safe_load(content)
        formatted = yaml.dump(data, allow_unicode=True, default_flow_style=False)
        return f"[YAML: {Path(path).name}]\n\n{formatted}"
    except ImportError:
        return read_txt(path)
    except Exception:
        return read_txt(path)


# ─── HTML olvasó ─────────────────────────────────────────────────────────────

def read_html(path: str | Path) -> str:
    try:
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.texts = []
                self.skip_tags = {"script", "style", "head", "meta", "link"}
                self.current_skip = False

            def handle_starttag(self, tag, attrs):
                if tag.lower() in self.skip_tags:
                    self.current_skip = True

            def handle_endtag(self, tag):
                if tag.lower() in self.skip_tags:
                    self.current_skip = False

            def handle_data(self, data):
                if not self.current_skip:
                    text = data.strip()
                    if text:
                        self.texts.append(text)

        content = Path(path).read_text(encoding="utf-8", errors="ignore")
        parser = TextExtractor()
        parser.feed(content)
        return f"[HTML: {Path(path).name}]\n\n" + "\n".join(parser.texts)
    except Exception:
        return read_txt(path)


# ─── Formátum → olvasó mapping ───────────────────────────────────────────────

READERS = {
    ".docx": read_docx,
    ".pdf":  read_pdf,
    ".xlsx": read_xlsx,
    ".xls":  read_xlsx,
    ".txt":  read_txt,
    ".md":   read_txt,
    ".log":  read_txt,
    ".py":   read_script,
    ".ps1":  read_script,
    ".sh":   read_script,
    ".bat":  read_script,
    ".cmd":  read_script,
    ".csv":  read_csv,
    ".xml":  read_xml,
    ".json": read_json,
    ".yaml": read_yaml,
    ".yml":  read_yaml,
    ".html": read_html,
    ".htm":  read_html,
}


def read_file(path: str | Path) -> str | None:
    """Visszaadja a fájl szöveges tartalmát, vagy None-t ha nem támogatott."""
    suffix = Path(path).suffix.lower()
    reader = READERS.get(suffix)
    if not reader:
        return None
    try:
        return reader(path)
    except Exception as e:
        print(f"[Olvasás hiba] {path}: {e}")
        return None


def supported_extensions() -> list[str]:
    """Visszaadja a támogatott kiterjesztések listáját."""
    return sorted(READERS.keys())
